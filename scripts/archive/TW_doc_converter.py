# /// script
# requires-python = ">=3.8"
# dependencies = [
#     "python-docx",
#     "markdown",
#     "beautifulsoup4",
#     "pillow",
# ]
# ///

"""
TW Document Converter
Student: TW
Description: Converting my writing sentences into doc documents

This script demonstrates how to:
1. Convert plain text to formatted Word documents
2. Apply various formatting styles (headings, bold, italic, etc.)
3. Add images, tables, and other elements
4. Parse markdown text and convert to Word
5. Batch convert multiple text files

Educational Focus:
- Working with python-docx library
- Document formatting and styling
- Text processing and parsing
- File I/O operations
- Object-oriented programming
"""

import os
import re
from pathlib import Path
from datetime import datetime
from typing import List, Dict, Optional, Union
import markdown
from bs4 import BeautifulSoup

try:
    from docx import Document
    from docx.shared import Inches, Pt
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.style import WD_STYLE_TYPE
    from docx.oxml.shared import OxmlElement, qn
except ImportError:
    print("❌ python-docx not installed. Run: pip install python-docx")
    exit(1)

class DocumentConverter:
    """
    A comprehensive text-to-Word document converter

    This class provides methods to:
    - Convert plain text to formatted Word documents
    - Parse and convert markdown to Word
    - Apply various formatting styles
    - Add headers, footers, and metadata
    - Batch process multiple files
    """

    def __init__(self):
        """Initialize the converter with default settings"""
        # Default formatting settings
        self.default_font = "Calibri"
        self.default_font_size = 11
        self.heading_font = "Calibri"

        # Style mappings for different text elements
        self.style_mappings = {
            'h1': 'Heading 1',
            'h2': 'Heading 2',
            'h3': 'Heading 3',
            'h4': 'Heading 4',
            'h5': 'Heading 5',
            'h6': 'Heading 6',
            'p': 'Normal',
            'blockquote': 'Quote',
            'code': 'Intense Quote'  # Using built-in style for code blocks
        }

        print("✅ Document Converter initialized")

    def create_document_from_text(self,
                                text: str,
                                title: str = "Document",
                                author: str = "TW",
                                apply_auto_formatting: bool = True) -> Document:
        """
        Create a Word document from plain text with automatic formatting

        Args:
            text: The text content to convert
            title: Document title
            author: Document author
            apply_auto_formatting: Whether to apply automatic formatting detection

        Returns:
            Document object ready to be saved
        """
        print(f"📝 Creating document: {title}")

        # Create new document
        doc = Document()

        # Set document properties
        doc.core_properties.title = title
        doc.core_properties.author = author
        doc.core_properties.created = datetime.now()

        # Add title
        title_paragraph = doc.add_heading(title, level=0)
        title_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER

        # Add author and date
        meta_paragraph = doc.add_paragraph()
        meta_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
        meta_run = meta_paragraph.add_run(f"By {author} | {datetime.now().strftime('%B %d, %Y')}")
        meta_run.italic = True

        # Add separator
        doc.add_paragraph()

        if apply_auto_formatting:
            self._add_formatted_text(doc, text)
        else:
            self._add_plain_text(doc, text)

        print(f"✅ Document created with {len(doc.paragraphs)} paragraphs")
        return doc

    def _add_formatted_text(self, doc: Document, text: str):
        """
        Add text with automatic formatting detection

        This method looks for common patterns in text and applies formatting:
        - Lines starting with # become headings
        - Lines in ALL CAPS become emphasized
        - Lines with bullet points become lists
        - Quoted text becomes block quotes
        """
        lines = text.split('\n')
        current_list = None

        for line in lines:
            line = line.strip()

            if not line:
                # Empty line - add paragraph break
                doc.add_paragraph()
                current_list = None
                continue

            # Check for heading patterns (# Header, ## Header, etc.)
            heading_match = re.match(r'^(#{1,6})\s+(.+)', line)
            if heading_match:
                level = len(heading_match.group(1))
                heading_text = heading_match.group(2)
                doc.add_heading(heading_text, level=level)
                current_list = None
                continue

            # Check for bullet points
            bullet_match = re.match(r'^[\-\*\+]\s+(.+)', line)
            if bullet_match:
                bullet_text = bullet_match.group(1)
                if current_list is None:
                    current_list = doc.add_paragraph(bullet_text, style='List Bullet')
                else:
                    doc.add_paragraph(bullet_text, style='List Bullet')
                continue

            # Check for numbered lists
            number_match = re.match(r'^\d+\.\s+(.+)', line)
            if number_match:
                number_text = number_match.group(1)
                doc.add_paragraph(number_text, style='List Number')
                current_list = None
                continue

            # Check for quotes (lines starting with >)
            quote_match = re.match(r'^>\s+(.+)', line)
            if quote_match:
                quote_text = quote_match.group(1)
                quote_paragraph = doc.add_paragraph(quote_text, style='Quote')
                current_list = None
                continue

            # Check for emphasis patterns
            if line.isupper() and len(line.split()) > 1:
                # ALL CAPS text becomes emphasized
                emphasis_paragraph = doc.add_paragraph()
                emphasis_run = emphasis_paragraph.add_run(line)
                emphasis_run.bold = True
                current_list = None
                continue

            # Regular paragraph
            paragraph = doc.add_paragraph(line)

            # Apply inline formatting (bold, italic)
            self._apply_inline_formatting(paragraph)
            current_list = None

    def _apply_inline_formatting(self, paragraph):
        """
        Apply inline formatting like **bold** and *italic* to paragraph text

        This method processes markdown-style formatting within paragraphs
        """
        # Get the text from the paragraph
        text = paragraph.text

        # Clear the paragraph
        paragraph.clear()

        # Process bold text (**text**)
        bold_pattern = r'\*\*(.*?)\*\*'
        parts = re.split(bold_pattern, text)

        for i, part in enumerate(parts):
            if i % 2 == 0:
                # Regular text, check for italic
                italic_pattern = r'\*(.*?)\*'
                italic_parts = re.split(italic_pattern, part)

                for j, italic_part in enumerate(italic_parts):
                    if j % 2 == 0:
                        # Regular text
                        if italic_part:
                            paragraph.add_run(italic_part)
                    else:
                        # Italic text
                        italic_run = paragraph.add_run(italic_part)
                        italic_run.italic = True
            else:
                # Bold text
                bold_run = paragraph.add_run(part)
                bold_run.bold = True

    def _add_plain_text(self, doc: Document, text: str):
        """Add text as plain paragraphs without formatting detection"""
        paragraphs = text.split('\n\n')  # Split on double newlines

        for para_text in paragraphs:
            if para_text.strip():
                doc.add_paragraph(para_text.strip())

    def convert_markdown_to_docx(self, markdown_text: str, title: str = "Markdown Document") -> Document:
        """
        Convert markdown text to a Word document

        Args:
            markdown_text: Markdown formatted text
            title: Document title

        Returns:
            Document object with converted content
        """
        print(f"🔄 Converting markdown to Word document: {title}")

        # Convert markdown to HTML
        html = markdown.markdown(markdown_text, extensions=['tables', 'fenced_code'])

        # Parse HTML with BeautifulSoup
        soup = BeautifulSoup(html, 'html.parser')

        # Create new document
        doc = Document()

        # Set document properties
        doc.core_properties.title = title
        doc.core_properties.author = "TW"
        doc.core_properties.created = datetime.now()

        # Add title
        doc.add_heading(title, level=0)

        # Process HTML elements
        for element in soup.find_all():
            self._process_html_element(doc, element)

        print(f"✅ Markdown converted to Word document")
        return doc

    def _process_html_element(self, doc: Document, element):
        """Process individual HTML elements and convert to Word format"""
        tag_name = element.name

        if tag_name in ['h1', 'h2', 'h3', 'h4', 'h5', 'h6']:
            # Headings
            level = int(tag_name[1])
            doc.add_heading(element.get_text().strip(), level=level)

        elif tag_name == 'p':
            # Paragraphs
            if element.get_text().strip():
                paragraph = doc.add_paragraph()
                self._add_formatted_text_to_paragraph(paragraph, element)

        elif tag_name == 'blockquote':
            # Block quotes
            quote_text = element.get_text().strip()
            if quote_text:
                doc.add_paragraph(quote_text, style='Quote')

        elif tag_name in ['ul', 'ol']:
            # Lists
            list_items = element.find_all('li')
            style = 'List Bullet' if tag_name == 'ul' else 'List Number'
            for item in list_items:
                doc.add_paragraph(item.get_text().strip(), style=style)

        elif tag_name == 'code':
            # Inline code - add as quote style
            code_text = element.get_text()
            if code_text.strip():
                doc.add_paragraph(code_text, style='Intense Quote')

    def _add_formatted_text_to_paragraph(self, paragraph, html_element):
        """Add formatted text from HTML element to a Word paragraph"""
        for content in html_element.contents:
            if hasattr(content, 'name'):
                # HTML element
                if content.name == 'strong' or content.name == 'b':
                    run = paragraph.add_run(content.get_text())
                    run.bold = True
                elif content.name == 'em' or content.name == 'i':
                    run = paragraph.add_run(content.get_text())
                    run.italic = True
                elif content.name == 'code':
                    run = paragraph.add_run(content.get_text())
                    run.font.name = 'Courier New'
                else:
                    paragraph.add_run(content.get_text())
            else:
                # Plain text
                paragraph.add_run(str(content))

    def add_table_from_data(self, doc: Document, data: List[List[str]], headers: List[str] = None):
        """
        Add a table to the document from data

        Args:
            doc: Document to add table to
            data: List of lists containing table data
            headers: Optional list of header names
        """
        if not data:
            return

        rows = len(data) + (1 if headers else 0)
        cols = len(data[0]) if data else 0

        # Create table
        table = doc.add_table(rows=rows, cols=cols)
        table.style = 'Table Grid'

        row_idx = 0

        # Add headers if provided
        if headers:
            header_row = table.rows[row_idx]
            for col_idx, header in enumerate(headers):
                header_row.cells[col_idx].text = header
                # Make header bold
                for paragraph in header_row.cells[col_idx].paragraphs:
                    for run in paragraph.runs:
                        run.bold = True
            row_idx += 1

        # Add data rows
        for data_row in data:
            table_row = table.rows[row_idx]
            for col_idx, cell_data in enumerate(data_row):
                table_row.cells[col_idx].text = str(cell_data)
            row_idx += 1

        print(f"✅ Added table with {rows} rows and {cols} columns")

    def batch_convert_files(self,
                          input_directory: Union[str, Path],
                          output_directory: Union[str, Path] = None,
                          file_pattern: str = "*.txt") -> List[Path]:
        """
        Batch convert multiple text files to Word documents

        Args:
            input_directory: Directory containing text files
            output_directory: Directory to save Word documents (defaults to input_directory)
            file_pattern: File pattern to match (e.g., "*.txt", "*.md")

        Returns:
            List of created document file paths
        """
        input_path = Path(input_directory)
        output_path = Path(output_directory) if output_directory else input_path

        if not input_path.exists():
            print(f"❌ Input directory does not exist: {input_path}")
            return []

        # Create output directory if it doesn't exist
        output_path.mkdir(parents=True, exist_ok=True)

        # Find matching files
        files = list(input_path.glob(file_pattern))
        created_docs = []

        print(f"🔄 Batch converting {len(files)} files from {input_path}")

        for file_path in files:
            try:
                # Read file content
                with open(file_path, 'r', encoding='utf-8') as f:
                    content = f.read()

                # Determine conversion method based on file extension
                if file_path.suffix.lower() == '.md':
                    doc = self.convert_markdown_to_docx(content, file_path.stem)
                else:
                    doc = self.create_document_from_text(content, file_path.stem)

                # Save document
                output_file = output_path / f"{file_path.stem}.docx"
                doc.save(str(output_file))
                created_docs.append(output_file)

                print(f"✅ Converted: {file_path.name} → {output_file.name}")

            except Exception as e:
                print(f"❌ Error converting {file_path.name}: {e}")

        print(f"🎉 Batch conversion complete: {len(created_docs)} documents created")
        return created_docs

    def save_document(self, doc: Document, filename: str, output_dir: Union[str, Path] = None) -> Path:
        """
        Save a document to the specified location

        Args:
            doc: Document to save
            filename: Name for the output file (with or without .docx extension)
            output_dir: Directory to save in (defaults to current script directory)

        Returns:
            Path to the saved file
        """
        if output_dir is None:
            output_dir = Path(__file__).parent

        output_path = Path(output_dir)
        output_path.mkdir(parents=True, exist_ok=True)

        # Ensure .docx extension
        if not filename.endswith('.docx'):
            filename += '.docx'

        file_path = output_path / filename
        doc.save(str(file_path))

        print(f"💾 Document saved: {file_path}")
        return file_path

def create_sample_content() -> Dict[str, str]:
    """Create sample content for demonstration"""
    return {
        'plain_text': """
Welcome to the Document Converter Demo

This is a sample document that demonstrates the capabilities of our text-to-Word converter.

MAIN FEATURES
The converter can handle various types of formatting automatically.

Key Features:
- Automatic heading detection
- Bold and italic text formatting
- Bullet point lists
- Numbered lists
- Block quotes

*This text will be italic* and **this text will be bold**.

# This is a Heading 1
## This is a Heading 2
### This is a Heading 3

> This is a block quote that will be formatted specially.

Regular paragraph text continues here with normal formatting.
        """,

        'markdown_text': """
# Markdown Document Example

This document demonstrates **markdown to Word** conversion capabilities.

## Features

### Text Formatting
- **Bold text**
- *Italic text*
- `Inline code`

### Lists
1. First numbered item
2. Second numbered item
3. Third numbered item

#### Bullet Points
- First bullet point
- Second bullet point
- Third bullet point

### Code Blocks
```python
def hello_world():
    print("Hello, World!")
```

### Quotes
> This is a blockquote in markdown
> that spans multiple lines

## Tables
| Name | Age | City |
|------|-----|------|
| Alice | 30 | New York |
| Bob | 25 | San Francisco |
| Charlie | 35 | Chicago |

## Conclusion
This demonstrates the power of automated document conversion!
        """
    }

def main():
    """
    Main function demonstrating the document converter functionality
    """
    print("📝 TW Document Converter - Educational Demo")
    print("=" * 50)

    # Initialize the converter
    converter = DocumentConverter()

    # Get sample content
    samples = create_sample_content()

    # Example 1: Convert plain text to Word
    print("\n📚 Example 1: Converting plain text to Word document...")
    plain_doc = converter.create_document_from_text(
        samples['plain_text'],
        title="Plain Text Conversion Demo",
        author="TW Student"
    )

    converter.save_document(plain_doc, "plain_text_demo")

    # Example 2: Convert markdown to Word
    print("\n📚 Example 2: Converting markdown to Word document...")
    markdown_doc = converter.convert_markdown_to_docx(
        samples['markdown_text'],
        title="Markdown Conversion Demo"
    )

    converter.save_document(markdown_doc, "markdown_demo")

    # Example 3: Create document with table
    print("\n📚 Example 3: Creating document with table...")
    table_doc = converter.create_document_from_text(
        "# Student Grades Report\n\nBelow are the current student grades:",
        title="Grades Report",
        author="TW Teacher"
    )

    # Add sample table data
    table_data = [
        ["Alice Johnson", "95", "A"],
        ["Bob Smith", "87", "B+"],
        ["Charlie Brown", "92", "A-"],
        ["Diana Prince", "98", "A+"]
    ]

    converter.add_table_from_data(
        table_doc,
        table_data,
        headers=["Student Name", "Score", "Grade"]
    )

    converter.save_document(table_doc, "grades_report")

    # Example 4: Demonstrate batch conversion (if sample files exist)
    print("\n📚 Example 4: Batch conversion demo...")

    # Create sample files for batch conversion
    sample_dir = Path(__file__).parent / "sample_texts"
    sample_dir.mkdir(exist_ok=True)

    sample_files = {
        "essay1.txt": "# My First Essay\n\nThis is the content of my first essay...",
        "notes.md": "## Class Notes\n\n- Important point 1\n- Important point 2",
        "story.txt": "# Short Story\n\nOnce upon a time, in a land far away..."
    }

    # Create sample files
    for filename, content in sample_files.items():
        sample_file = sample_dir / filename
        with open(sample_file, 'w', encoding='utf-8') as f:
            f.write(content)

    # Batch convert
    created_docs = converter.batch_convert_files(sample_dir)

    print(f"\n🎓 Educational Notes:")
    print("1. python-docx provides extensive formatting capabilities")
    print("2. Always handle file encoding properly (use UTF-8)")
    print("3. Consider memory usage when processing large documents")
    print("4. Validate input data before processing")
    print("5. Use appropriate exception handling for file operations")
    print(f"6. Created {len(created_docs)} documents in batch conversion")

    # Clean up sample files
    import shutil
    if sample_dir.exists():
        shutil.rmtree(sample_dir)
        print("🧹 Cleaned up sample files")

if __name__ == "__main__":
    main()